"""Word document parser — extracts paragraphs + tables."""
from pathlib import Path

from app.parsers.base import ParseResult


def parse_docx(path: str | Path) -> ParseResult:
    res = ParseResult(parser="docx")
    try:
        from docx import Document
    except ImportError as e:
        res.status = "error"
        res.error = f"python-docx not installed: {e}"
        return res

    try:
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs]
        tables = []
        for t in doc.tables:
            tables.append([[c.text for c in row.cells] for row in t.rows])
        res.data = {
            "text": "\n".join(paragraphs),
            "paragraphs": paragraphs,
            "tables": tables,
        }
    except Exception as e:
        res.status = "error"
        res.error = f"docx parse failed: {e}"
    return res
